#!/usr/bin/env python3
"""
Generator dokumen "Usulan Ide Tugas Akhir", direplikasi persis dari format
referensi (Usulan Ide_2305551013_Agus Arya Wiraguna.pdf):
- Page size Letter (8.5x11 in), margin kiri 4 cm, atas/kanan/bawah 3 cm
- Font Times New Roman 12pt, spasi 1.5, rata kiri-kanan (justify)
- Indentasi baris pertama paragraf 1.5 cm
- Tabel identitas 2 kolom (label bold | isi)
- Heading bernomor bold ("1.\tPermasalahan")
- Caption gambar "Gambar N " (bold) + deskripsi
- Blok tanda tangan Calon Pembimbing
- Daftar Pustaka gaya APA, hanging indent
"""

import copy
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

FONT = "Times New Roman"
SIZE = Pt(12)
SIZE_SMALL = Pt(10)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ---------------------------------------------------------------- helpers --


def set_cell_border(cell, **kwargs):
    """Apply borders to a single table cell. kwargs: top/bottom/left/right =
    dict(sz=int, val='single', color='000000')"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn(f"w:{key}"), str(val))


def all_borders(cell, sz=4, color="000000"):
    edge = dict(sz=sz, val="single", color=color)
    set_cell_border(cell, top=edge, left=edge, bottom=edge, right=edge)


def prevent_row_split(row):
    """Keep a table row intact instead of splitting it across a page break."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def repeat_row_as_header(row):
    """Repeat this row as a header on every page the table spans."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_fixed_layout(table, col_widths_in):
    """Force a fixed table layout and rewrite the column grid so the exact
    column widths are respected (Word/LibreOffice otherwise redistribute
    widths based on cell content)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w * 1440)))  # inches -> twips
        grid.append(col)
    tblPr.addnext(grid)


def set_font(run, bold=False, italic=False, size=SIZE):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)


def base_paragraph_format(
    p, justify=True, first_line_indent=True, line_spacing=1.5, space_after=0
):
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        pf.first_line_indent = Cm(1.5)
    else:
        pf.first_line_indent = None
    return p


def add_body_paragraph(doc, text, indent=True, bold_runs=None):
    """text may contain simple bold markers via list of (text, bold, italic) tuples
    if bold_runs is provided (list of tuples); else plain text."""
    p = doc.add_paragraph()
    base_paragraph_format(p, first_line_indent=indent)
    if bold_runs:
        for t, b, i in bold_runs:
            r = p.add_run(t)
            set_font(r, bold=b, italic=i)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_heading(doc, number, title):
    p = doc.add_paragraph()
    base_paragraph_format(p, first_line_indent=False)
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Cm(1.5))
    r = p.add_run(f"{number}.\t{title}")
    set_font(r, bold=True)
    return p


def add_numbered_item(doc, number, bold_runs):
    """Hanging-indent numbered paragraph, e.g. objectives list."""
    p = doc.add_paragraph()
    base_paragraph_format(p, first_line_indent=False)
    pf = p.paragraph_format
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-1.0)
    pf.tab_stops.add_tab_stop(Cm(1.0))
    r = p.add_run(f"{number}.\t")
    set_font(r)
    for t, b, i in bold_runs:
        r = p.add_run(t)
        set_font(r, bold=b, italic=i)
    return p


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def add_figure(doc, path, number, caption, width_in=5.74):
    """Insert the image framed inside a bordered 1x1 table (like the
    reference document), followed by a centered caption below the box."""
    from PIL import Image as _PILImage

    with _PILImage.open(path) as im:
        px_w, px_h = im.size
    aspect = px_h / px_w

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(width_in)
    all_borders(cell, sz=6, color="000000")
    set_cell_margins(cell, top=150, start=150, bottom=150, end=150)

    p = cell.paragraphs[0]
    base_paragraph_format(p, justify=False, first_line_indent=False, space_after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inner_width = width_in - 0.25
    run.add_picture(
        path, width=Inches(inner_width), height=Inches(inner_width * aspect)
    )

    cap = doc.add_paragraph()
    base_paragraph_format(cap, justify=False, first_line_indent=False, space_after=12)
    cap.paragraph_format.space_before = Pt(6)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = cap.add_run(f"Gambar {number} ")
    set_font(r1, bold=True, size=SIZE_SMALL)
    r2 = cap.add_run(caption)
    set_font(r2, bold=False, size=SIZE_SMALL)
    return cap


def add_simple_table(
    doc, headers, rows, col_widths_in=None, center_from=None, size=SIZE_SMALL
):
    """center_from: if set to a column index, body cells from that column
    onward are center-aligned (useful for a comparison/checklist matrix).
    size: font size for header and body cells (data tables use 10pt)."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        base_paragraph_format(p, justify=False, first_line_indent=False)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, bold=True, size=size)
        all_borders(hdr_cells[i])
    prevent_row_split(table.rows[0])
    repeat_row_as_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            base_paragraph_format(p, justify=False, first_line_indent=False)
            if center_from is not None and i >= center_from:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_font(r, size=size)
            all_borders(cells[i])
        prevent_row_split(table.rows[-1])
    if col_widths_in:
        for row in table.rows:
            for i, w in enumerate(col_widths_in):
                row.cells[i].width = Inches(w)
        set_fixed_layout(table, col_widths_in)
    return table


def add_reference(doc, text_runs):
    """text_runs: list of (text, italic) tuples for hanging-indent bibliography entry."""
    p = doc.add_paragraph()
    base_paragraph_format(p, first_line_indent=False, space_after=8)
    pf = p.paragraph_format
    pf.left_indent = Cm(1.5)
    pf.first_line_indent = Cm(-1.5)
    for t, i in text_runs:
        r = p.add_run(t)
        set_font(r, italic=i)
    return p


# ------------------------------------------------------------------ setup --

doc = Document()

# Normal style defaults
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = SIZE
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), FONT)

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Cm(4)
section.top_margin = Cm(3)
section.right_margin = Cm(3)
section.bottom_margin = Cm(3)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

TEXT_WIDTH_IN = 8.5 - (4 / 2.54) - (3 / 2.54)  # ~5.74 in

# ------------------------------------------------------------------ title --

title_p = doc.add_paragraph()
base_paragraph_format(title_p, justify=False, first_line_indent=False, space_after=12)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("Usulan Ide Tugas Akhir")
set_font(r, bold=True)

# ------------------------------------------------------------- identity table --

table = doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

col1_w = 2.36
col2_w = TEXT_WIDTH_IN - col1_w

JUDUL = (
    "Implementasi Sistem SOAR Open-Source Berbasis n8n untuk Deteksi dan "
    "Respons Ancaman Malware dan Phishing dengan Mitigasi Aktif "
    "Human-in-the-Loop"
)

rows_data = [
    ("NIM", [("2305551076", False)]),
    ("Nama", [("Ravi Arnan Irianto", False)]),
    ("Bidang Keahlian", [("Network and Cloud Computing", False)]),
    (
        "Topik Networking dan Cloud\nComputing",
        [("Network Security (Security Operations, Detection & Response)", False)],
    ),
    ("Judul", [(JUDUL, False)]),
    (
        "Dosen Pembimbing",
        [
            (
                "1. Ir. I Nyoman Piarsa, ST., MT., IPM.\n"
                "2. A.A. Kt Agung Cahyawan Wiranatha, S.T., M.T.",
                False,
            )
        ],
    ),
]

for label, value_parts in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].width = Inches(col1_w)
    row_cells[1].width = Inches(col2_w)

    p0 = row_cells[0].paragraphs[0]
    base_paragraph_format(p0, justify=False, first_line_indent=False)
    for j, line in enumerate(label.split("\n")):
        if j > 0:
            p0.add_run().add_break()
        r = p0.add_run(line)
        set_font(r, bold=True)
    all_borders(row_cells[0])

    p1 = row_cells[1].paragraphs[0]
    base_paragraph_format(p1, justify=True, first_line_indent=False)
    for text, bold in value_parts:
        for j, line in enumerate(text.split("\n")):
            if j > 0:
                p1.add_run().add_break()
            r = p1.add_run(line)
            set_font(r, bold=bold)
    all_borders(row_cells[1])

doc.add_paragraph().paragraph_format.space_after = Pt(0)

# =================================================================== 1 ====

add_heading(doc, 1, "Permasalahan")

add_body_paragraph(
    doc,
    "Security Operations Center (SOC) modern menghadapi tekanan operasional "
    "yang terus meningkat akibat volume alert yang sangat besar, keterbatasan "
    "jumlah analis, dan tuntutan waktu respons yang semakin singkat "
    "(Habibzadeh, Feyzi, & Ebrahimi Atani, 2026). Platform Security Information "
    "and Event Management (SIEM) komersial umumnya mampu mendeteksi ancaman "
    "secara real-time, tetapi masih memerlukan intervensi manual untuk "
    "menindaklanjuti setiap peringatan sehingga proses respons menjadi lambat "
    "dan rentan terhadap human error (Nadig & Vignesh, 2026). Di sisi lain, "
    "solusi Security Orchestration, Automation, and Response (SOAR) komersial "
    "seperti Splunk SOAR maupun Cortex XSOAR memiliki biaya lisensi yang "
    "tinggi sehingga sulit diadopsi oleh organisasi berskala kecil-menengah "
    "maupun lingkungan akademik (Nadig & Vignesh, 2026).",
)

add_body_paragraph(
    doc,
    "Kombinasi SIEM open-source Wazuh dengan mesin orkestrasi telah banyak "
    "diteliti sebagai alternatif SOAR berbiaya rendah. Namun demikian, "
    "implementasi yang ada umumnya memilih salah satu dari dua pendekatan "
    "yang sama-sama memiliki kelemahan. Pendekatan pertama menerapkan "
    "otomatisasi penuh tanpa keterlibatan manusia, seperti integrasi Wazuh "
    "dengan Shuffle untuk pemblokiran otomatis pada serangan application-layer "
    "(Thakker, More, & Kumar, 2025) maupun brute force (Farrel, 2024). "
    "Pendekatan ini berisiko menimbulkan false-positive yang dieksekusi "
    "secara \u201cbuta\u201d karena keputusan mitigasi tidak diverifikasi oleh analis. "
    "Pendekatan kedua bersifat satu arah, yaitu sistem hanya mengirimkan "
    "notifikasi ke analis melalui Telegram, Slack, atau email tanpa "
    "menyediakan mekanisme bagi analis untuk langsung menindaklanjuti "
    "peringatan tersebut dari kanal notifikasi yang sama (Farrel, 2024; "
    "Sarker, 2025; n8n, n.d.). Akibatnya, analis tetap harus berpindah ke "
    "dasbor SIEM secara manual untuk melakukan tindakan mitigasi, sehingga "
    "waktu tanggap (response time) tidak banyak berkurang.",
)

add_body_paragraph(
    doc,
    "Selain permasalahan pada sisi orkestrasi, ketergantungan terhadap "
    "otomatisasi yang tidak transparan (black-box) turut menurunkan "
    "kepercayaan analis terhadap keputusan sistem, sementara automasi yang "
    "gagal secara diam-diam (silent failure) menyebabkan penurunan cakupan "
    "deteksi tanpa disadari (Mohsin, Janicke, Ibrahim, Sarker, & Camtepe, "
    "2025). Studi mengenai kolaborasi manusia dan AI pada SOC menegaskan "
    "perlunya kerangka human-in-the-loop yang memetakan tingkat otonomi "
    "sistem terhadap tingkat keyakinan (confidence) dan risiko setiap "
    "keputusan, alih-alih menerapkan otomatisasi biner (sepenuhnya otomatis "
    "atau sepenuhnya manual) (Mohsin, et al., 2025). Pemanfaatan large "
    "language model (LLM) untuk membantu analisis maupun perencanaan respons "
    "insiden juga menunjukkan potensi besar, tetapi pemakaian LLM berbasis "
    "cloud menimbulkan risiko kedaulatan data serta biaya API yang berulang "
    "(Hammar, Alpcan, & Lupu, 2025).",
)

add_body_paragraph(
    doc,
    "Berdasarkan permasalahan tersebut, diperlukan sebuah sistem SOAR "
    "open-source yang mengintegrasikan Wazuh sebagai mesin deteksi (SIEM "
    "dan File Integrity Monitoring) dengan n8n sebagai mesin orkestrasi "
    "playbook, dilengkapi verifikasi ancaman multi-sumber (VirusTotal, "
    "Google Safe Browsing, dan URLScan.io), analisis AI on-premise, serta "
    "mekanisme mitigasi aktif (Active Response) yang bersifat human-in-the-"
    "loop dua arah melalui satu kanal chat yang sama. Sistem yang diusulkan "
    "diharapkan mampu memberikan respons berjenjang berdasarkan tingkat "
    "keyakinan ancaman malware dan phishing sehingga menekan false-positive "
    "dan mempercepat mitigasi tanpa mengorbankan kendali dan akuntabilitas "
    "keputusan analis manusia.",
)

# =================================================================== 2 ====

add_heading(doc, 2, "Tujuan Penelitian")

add_body_paragraph(
    doc,
    "Penelitian ini bertujuan untuk merancang dan mengimplementasikan "
    "Sistem SOAR Open-Source Berbasis n8n untuk Deteksi dan Respons Ancaman "
    "Malware dan Phishing dengan Mitigasi Aktif Human-in-the-Loop. Secara "
    "khusus, tujuan penelitian ini adalah sebagai berikut.",
)

objectives = [
    (
        "Merancang arsitektur sistem SOAR open-source yang mengintegrasikan "
        "Wazuh sebagai mesin deteksi (SIEM dan File Integrity Monitoring) "
        "dengan n8n sebagai mesin orkestrasi playbook untuk menangani dua "
        "kelas ancaman, yaitu malware dan phishing."
    ),
    (
        "Mengembangkan mekanisme respons berjenjang berbasis tingkat keyakinan "
        "(confidence-based tiered response) menggunakan verifikasi threat "
        "intelligence multi-sumber (VirusTotal, Google Safe Browsing, dan "
        "URLScan.io) agar tindakan mitigasi sebanding dengan tingkat keyakinan "
        "ancaman."
    ),
    (
        "Mengimplementasikan mekanisme mitigasi aktif (Active Response) "
        "human-in-the-loop dua arah melalui antarmuka Telegram, sehingga "
        "analis dapat menyetujui atau menolak tindakan karantina file maupun "
        "pemblokiran domain langsung dari kanal notifikasi."
    ),
    (
        "Mengintegrasikan analisis AI lokal (on-premise large language model) "
        "untuk menghasilkan ringkasan dan penjelasan (explainability) atas "
        "setiap keputusan deteksi tanpa mengirimkan data sensitif ke pihak "
        "ketiga."
    ),
    (
        "Menerapkan arsitektur multi-agent lintas distribusi (Ubuntu dan "
        "Rocky Linux) yang terhubung melalui mesh VPN, sehingga deteksi dan "
        "mitigasi tetap berjalan meski endpoint berpindah jaringan."
    ),
    (
        "Mengevaluasi kinerja sistem secara kuantitatif melalui pengukuran "
        "Mean Time To Respond (MTTR) dan tingkat reduksi false-positive pada "
        "skenario deteksi malware dan phishing."
    ),
]
for idx, obj in enumerate(objectives, start=1):
    add_numbered_item(doc, idx, [(obj, False, False)])

# =================================================================== 3 ====

add_heading(doc, 3, "Gambaran Umum Penelitian")

add_body_paragraph(
    doc,
    "Sistem yang akan dikembangkan merupakan Sistem SOAR (Security "
    "Orchestration, Automation, and Response) open-source yang berfungsi "
    "sebagai lapisan orkestrasi di atas SIEM Wazuh untuk menangani dua "
    "kelas ancaman keamanan, yaitu malware dan phishing. Wazuh berperan "
    "sebagai mesin deteksi melalui modul File Integrity Monitoring (FIM) "
    "dan korelasi log pada setiap endpoint, sedangkan n8n berperan sebagai "
    "\u201cotak\u201d SOAR yang mengeksekusi playbook otomatis mulai dari pengayaan "
    "data (enrichment), klasifikasi tingkat keyakinan, analisis AI, hingga "
    "eksekusi mitigasi (Alves, Loureiro, & Pedrosa, 2026). Pendekatan "
    "kombinasi SIEM dan mesin orkestrasi low-code seperti n8n dipilih "
    "karena terbukti dapat mereduksi beban triase manual dibandingkan pola "
    "SOC konvensional (Thakker, More, & Kumar, 2025), sekaligus tetap "
    "bersifat sepenuhnya open-source sehingga tidak menimbulkan biaya "
    "lisensi seperti pada SIEM/SOAR komersial (Nadig & Vignesh, 2026).",
)

add_body_paragraph(
    doc,
    "Arsitektur sistem terdiri atas lima komponen utama yang saling "
    "terintegrasi, yaitu Endpoint terdistribusi, SOAR Server terpusat, "
    "layanan Threat Intelligence eksternal, mesin AI lokal, dan kanal "
    "Human-in-the-Loop. Kelima komponen tersebut membentuk satu alur "
    "kerja SOAR yang utuh, yaitu",
)

comp_items = [
    (
        "Endpoint (terdistribusi) berupa Wazuh Agent yang berjalan pada "
        "beberapa distribusi Linux (Ubuntu dan Rocky Linux) untuk melakukan "
        "File Integrity Monitoring pada direktori sensitif serta pengumpulan "
        "log secara real-time. Setiap endpoint terhubung ke manager melalui "
        "mesh VPN Tailscale sehingga tetap dapat melapor walau berpindah "
        "jaringan."
    ),
    (
        "SOAR Server (terpusat) yang menjalankan Wazuh Manager beserta "
        "integratord sebagai pusat agregasi dan korelasi alert, serta n8n "
        "sebagai mesin orkestrasi yang mengeksekusi playbook Deteksi Malware, "
        "Deteksi Phishing, dan Telegram Callback Handler. Skrip jembatan "
        "custom-n8n.py meneruskan alert Wazuh ke webhook n8n setelah melalui "
        "penyaringan noise awal."
    ),
    (
        "Layanan Threat Intelligence eksternal, yaitu VirusTotal untuk "
        "reputasi hash file (malware) serta Google Safe Browsing dan "
        "URLScan.io untuk reputasi URL (phishing). Verifikasi multi-sumber "
        "ini menempatkan hasil korelasi Wazuh sebagai indikasi awal, bukan "
        "keputusan akhir, sehingga tindakan mitigasi hanya dieksekusi setelah "
        "dikonfirmasi oleh sumber reputasi eksternal."
    ),
    (
        "Mesin AI lokal (Ollama dengan model llama3.2:3b) yang menghasilkan "
        "ringkasan dan penjelasan keputusan (explainability) dalam Bahasa "
        "Indonesia tanpa mengirimkan data ke layanan pihak ketiga, sehingga "
        "kedaulatan data tetap terjaga dan tidak menimbulkan biaya API "
        "berulang (Hammar, Alpcan, & Lupu, 2025)."
    ),
    (
        "Kanal Human-in-the-Loop berupa Telegram Bot yang mengirimkan "
        "notifikasi dua arah dengan tombol interaktif (Isolasi/Blokir/Abaikan), "
        "serta tg-callback-poller yang melakukan long-poll terhadap Telegram "
        "API agar sistem tetap dapat menerima keputusan analis walau server "
        "berada di balik NAT tanpa perlu mengekspos webhook publik."
    ),
]
for item in comp_items:
    add_numbered_item(doc, comp_items.index(item) + 1, [(item, False, False)])

add_figure(
    doc,
    os.path.join(BASE_DIR, "assets/gambar-1-arsitektur-simplified.png"),
    1,
    "Arsitektur Sistem SOAR Open-Source Berbasis n8n.",
    width_in=TEXT_WIDTH_IN,
)

add_body_paragraph(
    doc,
    "Alur kerja sistem dimulai ketika Wazuh Agent mendeteksi perubahan "
    "berkas mencurigakan (malware) atau akses ke URL tertentu (phishing). "
    "Wazuh Manager mengorelasikan event tersebut menjadi alert yang "
    "diteruskan skrip custom-n8n.py ke webhook n8n. Playbook n8n kemudian "
    "melakukan verifikasi ke sumber threat intelligence terkait sebelum "
    "menetapkan tingkat keyakinan (severity) ancaman. Klasifikasi tingkat "
    "keyakinan untuk masing-masing kelas ancaman dirancang sebagai berikut.",
)

cap = doc.add_paragraph()
base_paragraph_format(cap, justify=False, first_line_indent=False, space_after=6)
cap.paragraph_format.keep_with_next = True
r1 = cap.add_run("Tabel 1 ")
set_font(r1, bold=True, size=SIZE_SMALL)
r2 = cap.add_run("Klasifikasi Tingkat Keyakinan Deteksi Malware.")
set_font(r2, size=SIZE_SMALL)

add_simple_table(
    doc,
    ["Hasil Verifikasi VirusTotal", "Tingkat Keyakinan", "Tindakan"],
    [
        [
            "malicious \u2265 20 mesin AV",
            "KRITIS",
            "Karantina file otomatis + notifikasi",
        ],
        [
            "malicious 1\u201319 mesin AV",
            "TINGGI",
            "Notifikasi + tombol Telegram (analis memutuskan)",
        ],
        [
            "malicious = 0 (bersih/tidak dikenal)",
            "AMAN",
            "Disenyapkan (tanpa notifikasi)",
        ],
    ],
    col_widths_in=[2.2, 1.3, 2.24],
)
doc.add_paragraph().paragraph_format.space_after = Pt(12)

cap = doc.add_paragraph()
base_paragraph_format(cap, justify=False, first_line_indent=False, space_after=6)
cap.paragraph_format.keep_with_next = True
r1 = cap.add_run("Tabel 2 ")
set_font(r1, bold=True, size=SIZE_SMALL)
r2 = cap.add_run("Klasifikasi Tingkat Keyakinan Deteksi Phishing.")
set_font(r2, size=SIZE_SMALL)

add_simple_table(
    doc,
    ["Hasil Verifikasi GSB dan URLScan", "Tingkat Keyakinan", "Tindakan"],
    [
        ["Terdeteksi berbahaya", "BERBAHAYA", "Sinkhole domain otomatis + notifikasi"],
        [
            "Skor mencurigakan",
            "MENCURIGAKAN",
            "Notifikasi + tombol Telegram (analis memutuskan)",
        ],
        [
            "Sumber tidak dapat diverifikasi (rate-limit)",
            "PERLU VERIFIKASI",
            "Eskalasi ke analis (tidak diklaim aman)",
        ],
        ["Bersih", "AMAN", "Tidak ada tindakan"],
    ],
    col_widths_in=[2.2, 1.3, 2.24],
)
doc.add_paragraph().paragraph_format.space_after = Pt(12)

add_body_paragraph(
    doc,
    "Untuk ancaman dengan tingkat keyakinan TINGGI atau MENCURIGAKAN, "
    "sistem tidak melakukan mitigasi secara otomatis, melainkan mengirim "
    "notifikasi Telegram yang dilengkapi dua tombol interaktif (misalnya "
    "\u201cIsolasi File\u201d dan \u201cAbaikan\u201d). Keputusan mitigasi baru dieksekusi "
    "setelah analis menekan salah satu tombol tersebut, sehingga tindakan "
    "Active Response tetap berada di bawah kendali manusia (human-in-the-"
    "loop), bukan otomatisasi buta (Mohsin, et al., 2025). Mekanisme ini "
    "menjadi pembeda utama dibandingkan implementasi sejenis yang bersifat "
    "sepenuhnya otomatis (Thakker, More, & Kumar, 2025; Farrel, 2024) "
    "maupun yang hanya mengirim notifikasi satu arah tanpa kanal keputusan "
    "interaktif (Sarker, 2025; n8n, n.d.).",
)

add_figure(
    doc,
    os.path.join(BASE_DIR, "assets/gambar-2-sequence-hitl.png"),
    2,
    "Sequence Diagram Mitigasi Aktif Human-in-the-Loop.",
    width_in=TEXT_WIDTH_IN,
)

add_body_paragraph(
    doc,
    "Diagram alur pada Gambar 2 menunjukkan mekanisme keputusan dua arah "
    "yang diusulkan. Ketika analis menekan tombol pada notifikasi Telegram, "
    "callback_query diteruskan oleh tg-callback-poller (mekanisme long-poll "
    "getUpdates) menuju webhook n8n lokal, tanpa perlu mengekspos webhook "
    "n8n ke jaringan publik. Playbook Telegram Callback Handler kemudian "
    "mem-parsing keputusan analis, melakukan autentikasi ke Wazuh API, dan "
    "memicu Active Response yang sesuai, yaitu quarantine-file untuk "
    "ancaman malware atau block-domain (sinkhole) untuk ancaman phishing. "
    "Setelah eksekusi selesai, pesan Telegram diperbarui untuk mencatat "
    "keputusan beserta identitas analis dan waktu eksekusi sebagai jejak "
    "audit (audit trail).",
)

add_body_paragraph(
    doc,
    "Sebagai bagian dari kesadaran keamanan platform SOAR itu sendiri, "
    "rancangan ini mempertimbangkan temuan kerentanan kritis pada n8n "
    "(CVE-2026-21858, dijuluki \u201cNi8mare\u201d) yang berpotensi disalahgunakan "
    "melalui webhook publik (Cisco Talos, 2026). Oleh karena itu, seluruh "
    "webhook n8n pada sistem yang diusulkan tidak diekspos ke internet "
    "publik; interaksi dengan Telegram dilakukan melalui pola long-poll "
    "keluar (outbound-only) yang aman di balik NAT, sehingga permukaan "
    "serangan (attack surface) tetap minimal.",
)

add_body_paragraph(
    doc,
    "Sebagai indikasi awal kelayakan pendekatan yang diusulkan, pengujian "
    "pendahuluan pada purwarupa sistem menunjukkan waktu tanggap "
    "(Mean Time To Respond) rata-rata sekitar 1,7 detik untuk isolasi "
    "malware otomatis dan sekitar 2,1 detik untuk pemblokiran domain "
    "phishing otomatis, serta reduksi false-positive hingga 100% pada "
    "pengujian berkas jinak dibandingkan baseline deteksi berbasis aturan "
    "saja. Angka tersebut akan diverifikasi lebih lanjut dengan jumlah "
    "sampel yang lebih besar pada tahap pengujian penelitian.",
)

add_body_paragraph(
    doc,
    "Dibandingkan penelitian maupun implementasi sejenis, kebaruan yang "
    "diusulkan bersifat integratif, yaitu menggabungkan (1) respons "
    "berjenjang berbasis keyakinan hasil verifikasi multi-sumber, "
    "(2) mitigasi aktif human-in-the-loop dua arah melalui satu kanal chat, "
    "(3) cakupan dua kelas ancaman sekaligus (malware dan phishing) dengan "
    "jenis Active Response berbeda, (4) analisis AI on-premise untuk "
    "menjaga kedaulatan data, dan (5) arsitektur multi-agent lintas "
    "distribusi Linux yang resilien terhadap perpindahan jaringan. "
    "Kombinasi kelima aspek tersebut belum ditemukan secara bersamaan pada "
    "penelitian maupun implementasi Wazuh berbasis SOAR yang ada saat ini "
    "(Thakker, More, & Kumar, 2025; Nadig & Vignesh, 2026; Alves, Loureiro, "
    "& Pedrosa, 2026). Untuk memperjelas hal tersebut, Tabel 3 merangkum "
    "fokus, metode, dan fitur respons penelitian serta implementasi "
    "terdahulu, sedangkan Tabel 4 memetakan kesenjangan (gap) masing-masing "
    "beserta kontribusi yang ditawarkan proyek ini untuk menutupnya.",
)

cap = doc.add_paragraph()
base_paragraph_format(cap, justify=False, first_line_indent=False, space_after=6)
cap.paragraph_format.keep_with_next = True
r1 = cap.add_run("Tabel 3 ")
set_font(r1, bold=True, size=SIZE_SMALL)
r2 = cap.add_run("Ringkasan Penelitian dan Implementasi Terdahulu.")
set_font(r2, size=SIZE_SMALL)

add_simple_table(
    doc,
    [
        "Sumber",
        "Fokus Utama",
        "Metode",
        "Fitur Respons dan Alert",
    ],
    [
        [
            "Thakker dkk. (2025): pertahanan otomatis serangan "
            "application-layer pada Windows (Wazuh dan Shuffle)",
            "Mitigasi otomatis serangan application-layer pada endpoint "
            "Windows",
            "Integrasi Wazuh dan Shuffle; pemblokiran otomatis berbasis "
            "aturan",
            "Active Response otomatis penuh; tanpa notifikasi interaktif",
        ],
        [
            "Farrel (2024): SIEM Wazuh dengan Active Response dan notifikasi "
            "Telegram untuk brute force",
            "Mitigasi serangan brute force pada sistem informasi",
            "Wazuh Active Response dengan notifikasi Telegram",
            "Active Response otomatis; notifikasi Telegram satu arah",
        ],
        [
            "Sarker (2025): otomasi SOC: Wazuh, n8n, VirusTotal, dan Gmail "
            "untuk deteksi file berbahaya",
            "Deteksi file berbahaya berbasis orkestrasi n8n",
            "Orkestrasi n8n dengan VirusTotal dan notifikasi Gmail",
            "Notifikasi email satu arah; tanpa mitigasi aktif",
        ],
        [
            "n8n (n.d.): template workflow: Wazuh ke VirusTotal dengan alert "
            "Slack",
            "Notifikasi deteksi file berbahaya ke Slack",
            "Template workflow n8n dengan VirusTotal",
            "Notifikasi Slack satu arah; tanpa mitigasi",
        ],
        [
            "Nadig & Vignesh (2026): peningkatan Wazuh SIEM melalui integrasi "
            "SOAR untuk respons otomatis",
            "Respons ancaman otomatis melalui integrasi SOAR",
            "Integrasi Wazuh dengan SOAR; respons otomatis",
            "Respons otomatis; tanpa keterlibatan analis",
        ],
        [
            "Alves dkk. (2026): Wazuh untuk deteksi insiden: perancangan dan "
            "implementasi",
            "Perancangan dan implementasi deteksi insiden berbasis Wazuh",
            "Deteksi dan korelasi insiden pada Wazuh",
            "Berorientasi deteksi; minim otomatisasi respons",
        ],
        [
            "Raed (2025): SOC bertenaga AI: Wazuh, n8n, Claude AI, dan MCP",
            "SOC end-to-end dari deteksi hingga respons dengan bantuan AI",
            "Orkestrasi n8n dengan LLM Claude berbasis cloud (via MCP)",
            "Respons dibantu AI; tanpa kanal keputusan dua arah",
        ],
        [
            "zhadyz (2025): AI_SOC: SOC open-source berbasis LLM multi-agent "
            "(Ollama, Wazuh, TheHive, RAG)",
            "SOC open-source dengan LLM multi-agent dan AI lokal",
            "LLM multi-agent, Ollama, Wazuh, dan TheHive",
            "Triase berbasis TheHive; tanpa mitigasi dua arah via chat",
        ],
    ],
    col_widths_in=[1.55, 1.35, 1.3, 1.54],
)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

cap = doc.add_paragraph()
base_paragraph_format(cap, justify=False, first_line_indent=False, space_after=6)
cap.paragraph_format.keep_with_next = True
r1 = cap.add_run("Tabel 4 ")
set_font(r1, bold=True, size=SIZE_SMALL)
r2 = cap.add_run("Kesenjangan Penelitian Terdahulu dan Kontribusi Proyek.")
set_font(r2, size=SIZE_SMALL)

add_simple_table(
    doc,
    [
        "Sumber",
        "Kesenjangan",
        "Proyek Saya",
    ],
    [
        [
            "Thakker dkk. (2025)",
            "Eksekusi tanpa verifikasi analis (risiko false-positive); hanya "
            "satu kelas ancaman",
            "Respons berjenjang berbasis verifikasi multi-sumber dan mitigasi "
            "human-in-the-loop dua arah sebelum eksekusi",
        ],
        [
            "Farrel (2024)",
            "Notifikasi satu arah, analis harus pindah ke dasbor; hanya brute "
            "force",
            "Notifikasi Telegram interaktif dua arah; analis menyetujui atau "
            "menolak mitigasi langsung dari chat",
        ],
        [
            "Sarker (2025)",
            "Tidak ada kanal keputusan interaktif; fokus malware saja; tanpa "
            "AI lokal",
            "Mitigasi aktif (karantina dan sinkhole) melalui keputusan analis "
            "serta analisis AI lokal",
        ],
        [
            "n8n (n.d.)",
            "Hanya notifikasi; malware saja; tanpa mitigasi dua arah maupun AI "
            "lokal",
            "Playbook lengkap dengan mitigasi aktif dua arah, dua kelas "
            "ancaman, dan AI lokal",
        ],
        [
            "Nadig & Vignesh (2026)",
            "Tanpa human-in-the-loop; belum berjenjang multi-sumber; tanpa AI "
            "lokal",
            "Respons berjenjang berbasis keyakinan dengan kendali analis dan "
            "analisis AI lokal",
        ],
        [
            "Alves dkk. (2026)",
            "Minim orkestrasi mitigasi, kanal human-in-the-loop, dan analisis "
            "AI",
            "Menambahkan lapisan orkestrasi SOAR (n8n) untuk mitigasi aktif "
            "dan analisis AI lokal",
        ],
        [
            "Raed (2025)",
            "AI berbasis cloud (risiko kedaulatan data dan biaya API); tanpa "
            "keputusan interaktif",
            "Analisis AI on-premise (Ollama) untuk kedaulatan data dan "
            "mitigasi human-in-the-loop dua arah",
        ],
        [
            "zhadyz (2025)",
            "Tanpa mitigasi aktif human-in-the-loop dua arah; belum berjenjang "
            "multi-sumber; multi-agent bukan lintas-distribusi endpoint",
            "Mitigasi aktif dua arah via Telegram, respons berjenjang "
            "multi-sumber, dan multi-agent lintas-distribusi endpoint",
        ],
    ],
    col_widths_in=[1.3, 2.22, 2.22],
)

# ------------------------------------------------------------ signature --

doc.add_paragraph().paragraph_format.space_after = Pt(0)
doc.add_paragraph().paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
base_paragraph_format(p, justify=False, first_line_indent=False)
r = p.add_run("Calon Pembimbing 1,")
set_font(r)

for _ in range(4):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
base_paragraph_format(p, justify=False, first_line_indent=False)
r = p.add_run("(Ir. I Nyoman Piarsa, ST., MT., IPM.)")
set_font(r)

doc.add_paragraph().paragraph_format.space_after = Pt(0)
doc.add_paragraph().paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
base_paragraph_format(p, justify=False, first_line_indent=False)
r = p.add_run("Calon Pembimbing 2,")
set_font(r)

for _ in range(4):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
base_paragraph_format(p, justify=False, first_line_indent=False)
r = p.add_run("(A.A. Kt Agung Cahyawan Wiranatha, S.T., M.T.)")
set_font(r)

doc.add_page_break()

# =================================================================== 4 ====

add_heading(doc, 4, "Daftar Pustaka")

references = [
    [
        (
            "Alves, D. P., Loureiro, J., & Pedrosa, T. (2026). Wazuh for incident detection: Design and implementation. In ",
            False,
        ),
        (
            "Advanced Research in Technologies, Information, Innovation and Sustainability",
            True,
        ),
        (
            " (ARTIIS 2025), Communications in Computer and Information Science, Vol. 2791. Springer. doi:https://doi.org/10.1007/978-3-032-16848-1_35",
            False,
        ),
    ],
    [
        ("Cisco Talos. (2026). ", False),
        (
            "The n8n n8mare: How threat actors are misusing AI workflow automation.",
            True,
        ),
        (
            " Cisco Talos Blog. https://blog.talosintelligence.com/the-n8n-n8mare/",
            False,
        ),
    ],
    [
        (
            "Farrel, F. I. (2024). Implementation of security information & event management (SIEM) Wazuh with active response and Telegram notification for mitigating brute force attacks on the GT-I2TI USAKTI information system. ",
            False,
        ),
        ("INTELMATICS, 4", True),
        ("(1). doi:https://doi.org/10.25105/ITM.V4I1.18529", False),
    ],
    [
        (
            "Habibzadeh, A., Feyzi, F., & Ebrahimi Atani, R. (2026). Large language models for security operations centers: A comprehensive survey. ",
            False,
        ),
        ("Journal of Electrical and Computer Engineering, 2026", True),
        (", 3383674. doi:https://doi.org/10.1155/jece/3383674", False),
    ],
    [
        (
            "Hammar, K., Alpcan, T., & Lupu, E. C. (2025). Incident response planning using a lightweight large language model with reduced hallucination. ",
            False,
        ),
        ("arXiv", True),
        (", 1\u201314. doi:https://doi.org/10.48550/arXiv.2508.05188", False),
    ],
    [
        ("isMalicious. (2024). ", False),
        (
            "SOC alert fatigue: How threat intelligence reduces false positives without hiding real attacks.",
            True,
        ),
        (
            " https://ismalicious.com/posts/soc-alert-fatigue-threat-intelligence-false-positives",
            False,
        ),
    ],
    [
        (
            "Mohsin, A., Janicke, H., Ibrahim, A., Sarker, I. H., & Camtepe, S. (2025). A unified framework for human-AI collaboration in security operations centers with trusted autonomy. ",
            False,
        ),
        ("arXiv", True),
        (", 1\u201320. doi:https://doi.org/10.48550/arXiv.2505.23397", False),
    ],
    [
        ("n8n. (n.d.). ", False),
        (
            "Malicious file detection & response: Wazuh to VirusTotal with Slack alerts",
            True,
        ),
        (
            " [Workflow template]. https://n8n.io/workflows/5997-malicious-file-detection-and-response-wazuh-to-virustotal-with-slack-alerts/",
            False,
        ),
    ],
    [
        (
            "Nadig, A. S., & Vignesh, C. (2026). Enhancing Wazuh SIEM capabilities through SOAR integration for automated threat response. ",
            False,
        ),
        (
            "International Journal of Engineering Research & Technology (IJERT), Proceedings of ACSCON, 14",
            True,
        ),
        ("(6).", False),
    ],
    [
        ("Raed, B. (2025). ", False),
        (
            "A complete AI-powered SOC: From detection to automated response with Wazuh, n8n, Claude AI, and MCP.",
            True,
        ),
        (
            " Medium. https://medium.com/@basharraed/a-complete-ai-powered-soc-from-detection-to-automated-response-with-wazuh-n8n-claude-ai-and-mcp-78e7fa80d986",
            False,
        ),
    ],
    [
        ("Sarker, B. (2025). ", False),
        (
            "SOC automation: Wazuh SIEM integration with n8n, VirusTotal, Gmail for malicious file detection.",
            True,
        ),
        (
            " Medium. https://medium.com/@bappesarker2010/soc-automation-forwazuh-siem-integration-with-n8n-virustotal-gmail-03f3ee7ef684",
            False,
        ),
    ],
    [
        (
            "Thakker, A., More, A., & Kumar, K. (2025). Automated defense against application-layer attacks on Windows systems using Wazuh and Shuffle. ",
            False,
        ),
        (
            "International Journal of Education, Science, Technology, and Engineering (IJESTE), 8",
            True,
        ),
        ("(1), 45\u201357. doi:10.36079/lamintang.ijeste-0801.842", False),
    ],
    [
        ("Wazuh. (2025). ", False),
        (
            "Wazuh and Shuffle announce technology partnership to deliver integrated security automation.",
            True,
        ),
        (
            " Wazuh Blog. https://wazuh.com/blog/wazuh-and-shuffle-announce-technology-partnership-to-deliver-integrated-security-automation/",
            False,
        ),
    ],
    [
        ("zhadyz. (2025). ", False),
        (
            "AI_SOC: Open-source AI-augmented SOC (LLM + multi-agent, Ollama, Wazuh, TheHive, RAG)",
            True,
        ),
        (" [Software repository]. GitHub. https://github.com/zhadyz/AI_SOC", False),
    ],
]

for ref in references:
    add_reference(doc, ref)

OUT = os.path.join(BASE_DIR, "Usulan Ide_2305551076_Ravi Arnan Irianto.docx")
doc.save(OUT)
print("Saved:", OUT)
